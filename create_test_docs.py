from docx import Document
from docx.shared import Pt
import os

os.makedirs("test_docs", exist_ok=True)

# Create template DOCX
doc = Document()
doc.add_heading('Introduction', level=1)
p = doc.add_paragraph('This is the introduction paragraph.')
p.style.font.size = Pt(12)

doc.add_heading('Methodology', level=1)
p2 = doc.add_paragraph('This is the methodology section.')
p2.style.font.size = Pt(12)

doc.add_heading('Conclusion', level=1)
p3 = doc.add_paragraph('This is the conclusion.')
p3.style.font.size = Pt(12)

doc.save('test_docs/template.docx')

# Create target DOCX missing Conclusion, has misordered section, and wrong font sizes
doc_target = Document()
doc_target.add_heading('Intro', level=1)  # Flexible naming
pt1 = doc_target.add_paragraph('This is the intro paragraph with wrong size.')
for run in pt1.runs:
    run.font.size = Pt(16)

# Misordered
doc_target.add_heading('Results', level=1) # Missing in template
pr = doc_target.add_paragraph('Some results.')

doc_target.add_heading('Methodology', level=1)
pt2 = doc_target.add_paragraph('This is methodology.')

# Missing Conclusion

doc_target.save('test_docs/target.docx')

print("Test docs created successfully.")
